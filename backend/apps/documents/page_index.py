"""
PageIndex wrapper for Solera document intelligence.
Reads each file type, extracts plain text, builds a PageIndex hierarchical tree,
and stores the result in DocumentTree.tree_json.
"""
import os
from django.conf import settings


def extract_text(file_path: str, file_type: str) -> str:
    """Extract full text from a file depending on its type."""
    abs_path = os.path.join(settings.MEDIA_ROOT, file_path)

    if file_type == "pdf":
        return _extract_pdf_text(abs_path)
    elif file_type == "pdf_image":
        return _extract_pdf_ocr(abs_path)
    elif file_type == "docx":
        return _extract_docx(abs_path)
    elif file_type == "xlsx":
        return _extract_xlsx(abs_path)
    elif file_type == "transcript":
        return _extract_transcript(abs_path)
    else:
        return _read_plain(abs_path)


def build_tree(file_id: str) -> dict:
    """
    Build a PageIndex hierarchical tree for a ClientFile.
    Returns the tree dict (also stored in DocumentTree).
    """
    from apps.documents.models import ClientFile, DocumentTree
    from django.utils import timezone

    cf = ClientFile.objects.select_related("client").get(pk=file_id)
    tree_obj, _ = DocumentTree.objects.get_or_create(
        file=cf,
        defaults={"client": cf.client, "build_status": "pending"},
    )
    tree_obj.build_status = "pending"
    tree_obj.save()

    try:
        text = extract_text(cf.file_path, cf.file_type)
        if not text.strip():
            raise ValueError("No text extracted from file")

        tree = _run_page_index(text, cf.name)

        tree_obj.tree_json = tree
        tree_obj.build_status = "complete"
        tree_obj.built_at = timezone.now()
        tree_obj.error_msg = ""
        tree_obj.save()

        # Store a short summary on the file itself
        root_summary = tree.get("summary", "") if isinstance(tree, dict) else ""
        if root_summary:
            cf.ai_summary = root_summary[:500]
        cf.ai_processed = True
        cf.save()

        return tree

    except Exception as e:
        tree_obj.build_status = "failed"
        tree_obj.error_msg = str(e)
        tree_obj.save()
        raise


def _run_page_index(text: str, doc_name: str) -> dict:
    """Call PageIndex library to build the hierarchical tree."""
    try:
        from page_index import PageIndex
        pi = PageIndex()
        tree = pi.build(text, title=doc_name)
        # PageIndex returns either a dict or a Pydantic model — normalize to dict
        if hasattr(tree, "dict"):
            return tree.dict()
        if hasattr(tree, "model_dump"):
            return tree.model_dump()
        return tree if isinstance(tree, dict) else {"title": doc_name, "summary": str(tree)[:500], "nodes": []}
    except ImportError:
        # PageIndex not installed — return a minimal stub tree
        return {
            "title": doc_name,
            "node_id": "0001",
            "start_index": 0,
            "end_index": max(0, text.count("\n")),
            "summary": text[:500],
            "nodes": [],
        }


# ── Extractors ─────────────────────────────────────────────────────────────────

def _extract_pdf_text(abs_path: str) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(abs_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    return "\n\n".join(pages)


def _extract_pdf_ocr(abs_path: str) -> str:
    """For image-based / scanned PDFs — try text extraction first, OCR as fallback."""
    import pdfplumber
    import pytesseract
    pages = []
    with pdfplumber.open(abs_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if not text.strip():
                try:
                    img = page.to_image(resolution=150).original
                    text = pytesseract.image_to_string(img)
                except Exception:
                    pass
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(abs_path: str) -> str:
    from docx import Document
    doc = Document(abs_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(abs_path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(abs_path, read_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            line = "\t".join(str(c) if c is not None else "" for c in row)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def _extract_transcript(abs_path: str) -> str:
    """Parse VTT or SRT into speaker-tagged plain text."""
    with open(abs_path, encoding="utf-8", errors="ignore") as f:
        raw = f.read()

    ext = os.path.splitext(abs_path)[1].lower()
    if ext == ".vtt":
        return _parse_vtt(raw)
    elif ext == ".srt":
        return _parse_srt(raw)
    return raw


def _parse_vtt(raw: str) -> str:
    import re
    lines = []
    for block in raw.split("\n\n"):
        block_lines = block.strip().splitlines()
        # Find speaker line: "Speaker Name: text" or just text line after timestamp
        for line in block_lines:
            if "-->" not in line and not line.strip().isdigit() and line.strip() and line != "WEBVTT":
                lines.append(line.strip())
    return "\n".join(lines)


def _parse_srt(raw: str) -> str:
    import re
    # Remove timestamps and sequence numbers
    cleaned = re.sub(r'\d+\n\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\n', '', raw)
    return cleaned.strip()


def _read_plain(abs_path: str) -> str:
    with open(abs_path, encoding="utf-8", errors="ignore") as f:
        return f.read()
